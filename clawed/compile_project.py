"""Project packet DOCX compiler for ProjectArc.

Compiles a ProjectArc into a student-facing project packet with day-by-day
roadmap, choice boards, graphic organizer, research databases, rubric,
debate prep sheet, and culminating performance instructions.
No LLM calls — pure mechanical compilation.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from clawed.models import ProjectArc

logger = logging.getLogger(__name__)


async def compile_project_packet(
    project: ProjectArc,
    output_dir: Path,
) -> Path:
    """Compile a student-facing project packet DOCX from a ProjectArc.

    Args:
        project: The ProjectArc source object.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    from clawed.export_theme import get_color_theme
    from clawed.io import safe_filename

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    theme = get_color_theme("history")
    primary_hex = theme["primary"]

    def _hex_rgb(h: str) -> RGBColor:
        try:
            return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        except (ValueError, IndexError):
            return RGBColor(0x8B, 0x69, 0x14)

    def _shaded_cell(cell, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(tc_pr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill_hex,
        }))

    def _para(text: str, bold: bool = False, italic: bool = False, size_pt: int = 11) -> None:
        run = doc.add_paragraph().add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"

    _project_title_page(doc, project, _hex_rgb, primary_hex)
    _project_roadmap(doc, project, _para)
    _project_choice_boards(doc, project, _para)
    _project_graphic_organizer(doc, project, _shaded_cell, primary_hex)
    _project_resource_library(doc, project, _para)
    _project_debate_prep(doc, project, _para)
    _project_rubric(doc, project, _shaded_cell, primary_hex)
    _project_culminating(doc, project, _para)

    safe = safe_filename(project.title)
    out_path = output_dir / f"{safe}_project_packet.docx"
    doc.save(str(out_path))
    logger.info("Project packet saved to %s", out_path)
    return out_path

def _project_title_page(doc, project, hex_rgb_fn, primary_hex):
    from docx.shared import Pt
    title_heading = doc.add_heading(project.title, level=0)
    for run in title_heading.runs:
        run.font.color.rgb = hex_rgb_fn(primary_hex)
    if project.essential_question:
        eq_run = doc.add_paragraph().add_run(f'Essential Question: "{project.essential_question}"')
        eq_run.bold = True
        eq_run.italic = True
        eq_run.font.size = Pt(14)
        eq_run.font.name = "Calibri"
    doc.add_paragraph("")
    doc.add_paragraph("Name: ___________________________  Date: ___________  Period: _____")
    doc.add_paragraph("")


def _project_roadmap(doc, project, _para):
    doc.add_heading("Day-by-Day Roadmap", level=1)
    _para("Track your progress! Check off each phase as you complete it.", italic=True)
    doc.add_paragraph("")
    for phase in project.phases:
        text = f"[ ] {phase.title}"
        if phase.objective:
            text += f" \u2014 {phase.objective}"
        _para(text, bold=True)
        if phase.student_deliverable:
            _para(f"    Deliverable: {phase.student_deliverable}")
        doc.add_paragraph("")
    doc.add_page_break()


def _project_choice_boards(doc, project, _para):
    if project.movement_options:
        doc.add_heading("Choose Your Topic", level=1)
        _para("Select ONE of the following topics to specialize in:")
        doc.add_paragraph("")
        for i, option in enumerate(project.movement_options, 1):
            _para(f"{i}. {option}")
        doc.add_paragraph("")
        _para("My Choice: _______________________________________", bold=True)
        doc.add_paragraph("")
    if project.format_options:
        doc.add_heading("Choose Your Project Format", level=1)
        _para("How will you show what you know? Choose ONE:")
        doc.add_paragraph("")
        for option in project.format_options:
            _para(f"\u2022 {option}")
        doc.add_paragraph("")
        _para("My Choice: _______________________________________", bold=True)
        doc.add_paragraph("")
    doc.add_page_break()


def _project_graphic_organizer(doc, project, _shaded_cell, primary_hex):
    from docx.shared import RGBColor
    if not project.graphic_organizer:
        return
    doc.add_heading("Research Notes", level=1)
    cols = [c.strip() for c in project.graphic_organizer.split("|") if c.strip()]
    if cols:
        tbl = doc.add_table(rows=6, cols=len(cols))
        tbl.style = "Table Grid"
        for j, col_name in enumerate(cols):
            cell = tbl.rows[0].cells[j]
            _shaded_cell(cell, primary_hex)
            cell.text = col_name
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Calibri"
    doc.add_paragraph("")


def _project_resource_library(doc, project, _para):
    if not project.resource_library:
        return
    doc.add_heading("Research Resource Library", level=1)
    _para("Skip the basic Google search! Use these trusted databases:", italic=True)
    doc.add_paragraph("")
    for resource in project.resource_library:
        _para(f"\u2022 {resource.get('name', '')}", bold=True)
        if resource.get("url"):
            _para(f"  {resource['url']}")
        if resource.get("description"):
            _para(f"  {resource['description']}", italic=True)
    doc.add_paragraph("")
    doc.add_page_break()


def _project_debate_prep(doc, project, _para):
    if not project.debate_prep_template:
        return
    doc.add_heading("Debate Prep Sheet", level=1)
    for line in project.debate_prep_template.split("\n"):
        if line.strip():
            _para(line)
        else:
            doc.add_paragraph("")
    doc.add_paragraph("")


def _project_rubric(doc, project, _shaded_cell, primary_hex):
    from docx.shared import RGBColor
    if not project.rubric_text:
        return
    doc.add_heading("Grading Rubric", level=1)
    raw_lines = project.rubric_text.strip().split("\n")
    lines = [ln for ln in raw_lines if ln.strip() and not ln.strip().startswith("|---")]
    if not lines:
        return
    rows_data = [[c.strip() for c in ln.split("|") if c.strip()] for ln in lines]
    rows_data = [r for r in rows_data if r]
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=num_cols)
    tbl.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = tbl.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    _shaded_cell(cell, primary_hex)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph("")


def _project_culminating(doc, project, _para):
    perf = project.culminating_performance
    if not perf or not perf.title:
        return
    doc.add_page_break()
    doc.add_heading(f"Culminating Event: {perf.title}", level=1)
    if perf.format:
        _para(f"Format: {perf.format.replace('_', ' ').title()}  |  Duration: {perf.duration_minutes} minutes")
    if perf.setup_instructions:
        _para("Setup:", bold=True)
        _para(perf.setup_instructions)
    if perf.student_prep:
        _para("What You Need Ready:", bold=True)
        _para(perf.student_prep)
    if perf.evaluation_criteria:
        _para("You Will Be Evaluated On:", bold=True)
        for criterion in perf.evaluation_criteria:
            doc.add_paragraph(style="List Bullet").add_run(criterion)

