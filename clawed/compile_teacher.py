"""Teacher-view DOCX compiler for MasterContent.

Compiles a MasterContent object into a teacher-facing Word document with
full answer keys, teacher scripts, and all instructional notes.
No LLM calls — pure mechanical compilation.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.shared import RGBColor

    from clawed.master_content import MasterContent

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# Internal helpers (shared across section builders)
# ═══════════════════════════════════════════════════════════════════════════


class _DocBuilder:
    """Thin wrapper holding a python-docx Document and theme state."""

    def __init__(self, doc: DocxDocument, primary_hex: str,
                 accent_hex: str, bg_light_hex: str,
                 images: dict[str, Path]) -> None:
        from docx.shared import RGBColor
        self.doc = doc
        self.primary_hex = primary_hex
        self.accent_hex = accent_hex
        self.bg_light_hex = bg_light_hex
        self.images = images
        self._RGBColor = RGBColor

    def hex_rgb(self, h: str) -> RGBColor:
        return self._RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def heading(self, text: str, level: int = 1) -> None:
        h = self.doc.add_heading(text, level=level)
        if level <= 2:
            for run in h.runs:
                run.font.color.rgb = self.hex_rgb(self.primary_hex)

    def para(self, text: str, bold: bool = False, italic: bool = False,
             size_pt: int = 11, color: tuple[int, int, int] | None = None) -> None:
        from docx.shared import Pt
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = self._RGBColor(*color)

    def page_break(self) -> None:
        self.doc.add_page_break()  # type: ignore[no-untyped-call]  # python-docx is untyped

    def embed_image(self, image_spec: str, width_inches: float = 4.5) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        if not image_spec:
            return
        path = self.images.get(image_spec)
        if path and Path(path).exists():
            try:
                self.doc.add_picture(str(path), width=Inches(width_inches))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                logger.debug("Could not embed image %r: %s", image_spec, exc)

    def shaded_cell(self, cell: Any, fill_hex: str) -> None:
        from docx.oxml.ns import qn
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()  # noqa: N806
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill_hex,
        })
        tcPr.append(shd)

    def callout_box(self, title: str, items: list[str], fill_hex: str,
                    title_hex: str = "FFFFFF") -> None:
        """Render a colored callout box as a single-column table."""
        from docx.shared import Pt
        tbl = self.doc.add_table(rows=1 + len(items), cols=1)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells[0]
        self.shaded_cell(hdr, fill_hex)
        run = hdr.paragraphs[0].add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.hex_rgb(title_hex)
        run.font.name = "Calibri"
        for i, item in enumerate(items):
            cell = tbl.rows[i + 1].cells[0]
            self.shaded_cell(cell, self.bg_light_hex)
            run = cell.paragraphs[0].add_run(f"\u2022 {item}")
            run.font.size = Pt(10)
            run.font.name = "Calibri"
        self.doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════════════════
# Section builders
# ═══════════════════════════════════════════════════════════════════════════


def _build_header(b: _DocBuilder, master: MasterContent) -> None:
    """Title, metadata, and Materials-at-a-Glance table."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    b.doc.add_heading(master.title, level=0)

    meta_lines = [
        f"Subject: {master.subject}  |  Grade: {master.grade_level}  |  "
        f"Duration: {master.duration_minutes} min",
        f"Topic: {master.topic}",
        f"Objective: {master.objective}",
    ]
    for line in meta_lines:
        b.para(line)

    if master.standards:
        b.para("Standards: " + ", ".join(master.standards), bold=True)
    if master.materials_needed:
        b.para("Materials: " + ", ".join(master.materials_needed))
    b.doc.add_paragraph("")

    # Materials at a Glance table
    glance_format = getattr(master, "lesson_format", None) or "document_analysis"
    glance_format_display = glance_format.replace("_", " ").title()
    glance_standards = ", ".join(master.standards) if master.standards else "N/A"
    glance_ps_count = len(master.primary_sources) if master.primary_sources else 0
    glance_vocab_count = len(master.vocabulary) if master.vocabulary else 0
    glance_et_count = len(master.exit_ticket) if master.exit_ticket else 0

    diff = master.differentiation
    glance_iep = "IEP \u2713" if diff.struggling else "IEP \u2717"
    glance_ell = "ELL \u2713" if diff.ell else "ELL \u2717"
    glance_gifted = "Gifted \u2713" if diff.advanced else "Gifted \u2717"

    glance_table = b.doc.add_table(rows=4, cols=1)
    glance_table.style = "Table Grid"

    header_cell = glance_table.rows[0].cells[0]
    b.shaded_cell(header_cell, b.primary_hex)
    header_run = header_cell.paragraphs[0].add_run("MATERIALS AT A GLANCE")
    header_run.bold = True
    header_run.font.size = Pt(13)
    header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_run.font.name = "Calibri"
    header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        (f"Duration: {master.duration_minutes} min  |  "
         f"Format: {glance_format_display}  |  "
         f"Standards: {glance_standards}"),
        (f"Primary Sources: {glance_ps_count}  |  "
         f"Vocabulary Terms: {glance_vocab_count}  |  "
         f"Exit Ticket: {glance_et_count} questions"),
        (f"Differentiation: {glance_iep}  |  {glance_ell}  |  {glance_gifted}"),
    ]
    for idx, text in enumerate(rows_data, start=1):
        cell = glance_table.rows[idx].cells[0]
        b.shaded_cell(cell, b.accent_hex)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    b.doc.add_paragraph("")


def _build_vocabulary(b: _DocBuilder, master: MasterContent) -> None:
    """Vocabulary table with definitions and context sentences."""
    from docx.shared import RGBColor
    if not master.vocabulary:
        return
    b.heading("Vocabulary")
    table = b.doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ["Term", "Definition", "Context Sentence"], strict=False):
        b.shaded_cell(cell, b.primary_hex)
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for entry in master.vocabulary:
        row = table.add_row().cells
        row[0].text = entry.term
        row[1].text = entry.definition
        row[2].text = entry.context_sentence
        if entry.image_spec:
            b.embed_image(entry.image_spec, width_inches=1.5)

    b.doc.add_paragraph("")


def _build_do_now(b: _DocBuilder, master: MasterContent) -> None:
    """Do Now section with answer key."""
    b.heading("Do Now")
    b.para(master.do_now.stimulus)
    if master.do_now.questions:
        for i, q in enumerate(master.do_now.questions, 1):
            b.para(f"{i}. {q}")
    if master.do_now.answers:
        b.para("ANSWERS:", bold=True, color=(0x00, 0x70, 0xC0))
        for i, ans in enumerate(master.do_now.answers, 1):
            b.para(f"{i}. {ans}", italic=True)
    b.doc.add_paragraph("")


def _build_direct_instruction(b: _DocBuilder, master: MasterContent) -> None:
    """Direct instruction sections with teacher scripts."""
    b.heading("Direct Instruction")
    for section in master.direct_instruction:
        b.heading(section.heading, level=2)
        b.para(section.content)
        if section.key_points:
            b.para("Key Points:", bold=True)
            for kp in section.key_points:
                p = b.doc.add_paragraph(style="List Bullet")
                p.add_run(kp)
        if section.teacher_script:
            b.para("Teacher Script:", bold=True, color=(0x70, 0x30, 0xA0))
            b.para(section.teacher_script, italic=True, color=(0x70, 0x30, 0xA0))
        b.embed_image(section.image_spec)
        b.doc.add_paragraph("")


def _build_guided_notes(b: _DocBuilder, master: MasterContent) -> None:
    """Guided notes with answers filled in."""
    if not master.guided_notes:
        return
    b.heading("Guided Notes (Answer Key)")
    for note in master.guided_notes:
        b.para(f"Prompt: {note.prompt}", bold=True)
        b.para(f"Answer: {note.answer}", italic=True, color=(0x00, 0x70, 0xC0))
        if note.section_ref:
            b.para(f"(Ref: {note.section_ref})", size_pt=9, color=(0x66, 0x66, 0x66))
        b.doc.add_paragraph("")


def _build_primary_sources(b: _DocBuilder, master: MasterContent) -> None:
    """Primary source documents with scaffolding questions."""
    if not master.primary_sources:
        return
    b.page_break()
    b.heading("Primary Sources")
    for ps in master.primary_sources:
        b.heading(ps.title, level=2)
        b.para(f"Type: {ps.source_type}  |  Attribution: {ps.attribution}")
        b.para(ps.content_text)
        if ps.scaffolding_questions:
            b.para("Scaffolding Questions:", bold=True)
            for sq in ps.scaffolding_questions:
                p = b.doc.add_paragraph(style="List Bullet")
                p.add_run(sq)
        b.embed_image(ps.image_spec)
        b.doc.add_paragraph("")


def _build_activities(b: _DocBuilder, master: MasterContent) -> None:
    """Stations, jigsaw structure, and creative activities."""
    # Stations
    if master.stations:
        b.heading("Learning Stations")
        for station in master.stations:
            b.heading(station.title, level=2)
            b.para(f"Source: {station.source_ref}")
            b.para(f"Task: {station.task}")
            b.para("Student Directions:", bold=True)
            b.para(station.student_directions)
            b.para("Answer Key:", bold=True, color=(0xC0, 0x00, 0x00))
            b.para(station.teacher_answer_key, italic=True, color=(0xC0, 0x00, 0x00))
            b.doc.add_paragraph("")

    # Jigsaw
    jigsaw = getattr(master, "jigsaw", None)
    if jigsaw:
        b.heading("Jigsaw Activity Structure")
        b.para(f"Expert Groups: {jigsaw.num_expert_groups} groups", bold=True)
        b.para(f"Expert Phase: {jigsaw.expert_phase_minutes} minutes  |  "
               f"Teaching Phase: {jigsaw.teaching_phase_minutes} minutes")
        if jigsaw.documents_per_group:
            b.para("Documents per group: " + ", ".join(jigsaw.documents_per_group))
        if jigsaw.share_out_protocol:
            b.para("Share-Out Protocol:", bold=True)
            b.para(jigsaw.share_out_protocol)
        if jigsaw.graphic_organizer:
            b.para("Graphic Organizer Columns:", bold=True)
            b.para(jigsaw.graphic_organizer)
        if jigsaw.debrief_question:
            b.para("Debrief Question:", bold=True)
            b.para(jigsaw.debrief_question)
        b.doc.add_paragraph("")

    # Creative activity
    creative = getattr(master, "creative_activity", None)
    if creative and creative.title:
        b.heading(f"Creative Activity: {creative.title}")
        if creative.activity_type:
            b.para(f"Type: {creative.activity_type.replace('_', ' ').title()}  |  "
                   f"Time: {creative.time_minutes} minutes")
        if creative.scenario:
            b.para("Scenario:", bold=True)
            b.para(creative.scenario)
        if creative.roles:
            b.para("Roles:", bold=True)
            for role in creative.roles:
                p = b.doc.add_paragraph(style="List Bullet")
                p.add_run(role)
        if creative.student_directions:
            b.para("Student Directions:", bold=True)
            b.para(creative.student_directions)
        if creative.deliverable:
            b.para(f"Deliverable: {creative.deliverable}", bold=True)
        if creative.debrief:
            b.para("Debrief:", bold=True)
            b.para(creative.debrief)
        b.doc.add_paragraph("")


def _build_assessment(b: _DocBuilder, master: MasterContent) -> None:
    """Exit ticket with answer key."""
    if not master.exit_ticket:
        return
    b.page_break()
    b.heading("Exit Ticket")
    for i, sq in enumerate(master.exit_ticket, 1):
        b.para(f"Q{i} Stimulus ({sq.stimulus_type}): {sq.stimulus}", bold=True)
        if sq.stimulus_image_spec:
            b.embed_image(sq.stimulus_image_spec)
        b.para(f"Question: {sq.question}")
        starters = getattr(sq, "sentence_starters", [])
        if starters:
            b.para("Sentence Starters:", bold=True, color=(0x00, 0x70, 0xC0))
            for starter in starters:
                b.para(f"  \u2022 {starter}", color=(0x00, 0x70, 0xC0))
        b.para(f"Expected Answer: {sq.answer}", italic=True, color=(0xC0, 0x00, 0x00))
        if sq.cognitive_level:
            b.para(f"Cognitive Level: {sq.cognitive_level}", size_pt=9,
                   color=(0x66, 0x66, 0x66))
        b.doc.add_paragraph("")


def _build_differentiation(b: _DocBuilder, master: MasterContent) -> None:
    """Differentiation callout boxes."""
    diff = master.differentiation
    if diff.struggling or diff.advanced or diff.ell:
        b.page_break()
    b.heading("Differentiation")
    if diff.struggling:
        b.callout_box(
            "\u2691 IEP / 504 Accommodations",
            diff.struggling,
            fill_hex="D4A017",
            title_hex="FFFFFF",
        )
    if diff.advanced:
        b.callout_box(
            "\u2605 Advanced / Gifted Extensions",
            diff.advanced,
            fill_hex="2B7A98",
            title_hex="FFFFFF",
        )
    if diff.ell:
        b.callout_box(
            "\u2709 ELL Language Supports",
            diff.ell,
            fill_hex="2D8B4E",
            title_hex="FFFFFF",
        )


def _build_closure(b: _DocBuilder, master: MasterContent) -> None:
    """Independent work, homework, misconceptions, formative checks, prereqs."""
    # Independent work
    if master.independent_work:
        b.heading("Independent Work")
        b.para(master.independent_work.task)
        if master.independent_work.rubric_snippet:
            b.para("Rubric:", bold=True)
            b.para(master.independent_work.rubric_snippet)
        if master.independent_work.exemplar:
            b.para("Exemplar:", bold=True)
            b.para(master.independent_work.exemplar)
        b.doc.add_paragraph("")

    # Homework
    if master.homework:
        b.heading("Homework")
        b.para(master.homework)

    # Misconceptions
    if getattr(master, "misconceptions", None):
        b.heading("Common Misconceptions to Watch For")
        b.para(
            "Students often have these misunderstandings about this topic. "
            "Listen for them during discussion and address directly.",
            size_pt=10, italic=True, color=(0x66, 0x66, 0x66),
        )
        for item in master.misconceptions:
            p = b.doc.add_paragraph(style="List Bullet")
            p.add_run(item)
        b.doc.add_paragraph("")

    # Formative checks
    if getattr(master, "formative_checks", None):
        b.heading("Mid-Lesson Check-for-Understanding")
        b.para(
            "Ask these during instruction — verbal or show-of-hands. "
            "Catch confusion before it compounds.",
            size_pt=10, italic=True, color=(0x66, 0x66, 0x66),
        )
        for item in master.formative_checks:
            p = b.doc.add_paragraph(style="List Bullet")
            p.add_run(item)
        b.doc.add_paragraph("")

    # Prerequisites
    if getattr(master, "prerequisite_skills", None):
        b.heading("Prerequisite Skills")
        b.para(
            "Students should have these before this lesson. "
            "If not, plan a quick review.",
            size_pt=10, italic=True, color=(0x66, 0x66, 0x66),
        )
        for item in master.prerequisite_skills:
            p = b.doc.add_paragraph(style="List Bullet")
            p.add_run(item)
        b.doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════════════════
# Public API (unchanged signature)
# ═══════════════════════════════════════════════════════════════════════════


async def compile_teacher_view(
    master: MasterContent,
    images: dict[str, Path],
    output_dir: Path,
) -> Path:
    """Compile a teacher-facing DOCX from a MasterContent object.

    Includes full answer keys, teacher scripts (italicised), guided notes
    with answers filled in, station answer keys, and differentiation notes.

    Args:
        master: The MasterContent source-of-truth object.
        images: Mapping of image_spec strings to local file Paths.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document

    from clawed.export_theme import get_color_theme
    from clawed.io import safe_filename

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Subject-aware color theme
    theme = get_color_theme(master.subject)
    b = _DocBuilder(
        doc=doc,
        primary_hex=theme["primary"],
        accent_hex=theme["accent"],
        bg_light_hex=theme["bg_light"],
        images=images,
    )

    # Build all sections
    _build_header(b, master)
    _build_vocabulary(b, master)
    _build_do_now(b, master)
    _build_direct_instruction(b, master)
    _build_guided_notes(b, master)
    _build_primary_sources(b, master)
    _build_activities(b, master)
    _build_assessment(b, master)
    _build_differentiation(b, master)
    _build_closure(b, master)

    # Save
    safe = safe_filename(master.title)
    out_path = output_dir / f"{safe}_teacher.docx"
    doc.save(str(out_path))
    logger.info("Teacher view saved to %s", out_path)
    return out_path
