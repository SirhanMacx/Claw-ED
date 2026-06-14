"""Tool: generate_lesson_bundle — complete teaching package (lesson + handout + slides)."""
from __future__ import annotations

import asyncio
import logging
from pathlib import Path
from typing import Any

from clawed.agent_core.context import AgentContext, ToolResult
from clawed.failure_codes import FailureCode

logger = logging.getLogger(__name__)

_CJK_RANGES = [
    (0x4E00, 0x9FFF),    # CJK Unified Ideographs
    (0x3400, 0x4DBF),    # CJK Extension A
    (0x3000, 0x303F),    # CJK Symbols
    (0xFF00, 0xFFEF),    # Fullwidth Forms
    (0xAC00, 0xD7AF),    # Korean Hangul
]


def _has_cjk(text: str) -> bool:
    """Check if text contains CJK characters."""
    return any(
        any(lo <= ord(c) <= hi for lo, hi in _CJK_RANGES)
        for c in text
    )


def _strip_cjk(text: str) -> str:
    """Remove CJK characters from text, preserving everything else."""
    import re
    # Remove CJK blocks and clean up resulting whitespace
    cleaned = re.sub(
        r"[\u3000-\u303F\u3400-\u4DBF\u4E00-\u9FFF"
        r"\uAC00-\uD7AF\uFF00-\uFFEF]+",
        " ", text,
    )
    return re.sub(r"  +", " ", cleaned).strip()


def _sanitize_master_content(master: Any) -> None:
    """Strip CJK characters from all text fields in MasterContent.

    The minimax model sometimes outputs Chinese when generating
    historical quotes. This sanitizer catches it post-generation.
    """
    for field_name in dir(master):
        if field_name.startswith("_"):
            continue
        val = getattr(master, field_name, None)
        if isinstance(val, str) and _has_cjk(val):
            setattr(master, field_name, _strip_cjk(val))
        elif isinstance(val, list):
            for item in val:
                if hasattr(item, "__dict__"):
                    for k, v in vars(item).items():
                        if isinstance(v, str) and _has_cjk(v):
                            setattr(item, k, _strip_cjk(v))


async def _search_teacher_materials(
    context: AgentContext, topic: str, report: Any,
) -> tuple[str, str]:
    """Search for teacher's existing materials (assets + KB + KG).

    Returns (kb_context, kb_prompt_section).
    """
    kb_context = ""
    kb_prompt_section = ""

    # Asset-level search
    try:
        from clawed.asset_registry import AssetRegistry
        registry = AssetRegistry()
        assets = registry.search_assets(context.teacher_id, topic, top_k=5)
        yt_links = registry.get_youtube_links(context.teacher_id, topic, top_k=3)
        if assets or yt_links:
            kb_prompt_section = registry.format_asset_summary(assets, yt_links)
            logger.info("Asset search found %d files, %d YouTube links for '%s'",
                        len(assets), len(yt_links), topic)
    except Exception as e:
        logger.warning("NLAH_FAILURE=%s: %s", FailureCode.ASSET_SEARCH_FAILED, e)
        report.warnings.append(f"[{FailureCode.ASSET_SEARCH_FAILED}] Asset search failed: {e}")

    # KB chunk-level search
    try:
        from clawed.agent_core.memory.curriculum_kb import CurriculumKB
        kb = CurriculumKB()
        kb_results = kb.search(context.teacher_id, topic, top_k=3)
        if kb_results:
            kb_parts = [r for r in kb_results if r.get("similarity", 0) > 0.1]
            if kb_parts:
                kb_context = (
                    "\n\nRelevant materials from the teacher's files:\n"
                    + "\n".join(f"From '{r['doc_title']}': {r['chunk_text'][:200]}" for r in kb_parts)
                )
                chunk_section = "\n\n".join(
                    f"From \"{r['doc_title']}\":\n{r['chunk_text'][:500]}" for r in kb_parts
                )
                if kb_prompt_section:
                    kb_prompt_section += "\n\n" + chunk_section
                else:
                    kb_prompt_section = (
                        "Teacher's Existing Materials on This Topic\n"
                        "The teacher has created content on this topic before. "
                        "Reference and build on their existing work:\n\n"
                        + chunk_section
                        + "\n\nUse these materials as a foundation. "
                        "Reference the teacher's existing lessons, reuse their "
                        "graphic organizer formats, build on their approach."
                    )
                logger.info("KB search found %d relevant chunks for '%s'", len(kb_parts), topic)
    except Exception as e:
        logger.warning("NLAH_FAILURE=%s: %s", FailureCode.KB_SEARCH_FAILED, e)
        report.warnings.append(f"[{FailureCode.KB_SEARCH_FAILED}] KB search failed: {e}")

    # Knowledge graph context
    try:
        from clawed.agent_core.memory.knowledge_graph import CurriculumKG
        kg = CurriculumKG()
        kg_topic_context = kg.get_topic_context(context.teacher_id, topic)
        if kg_topic_context:
            if kb_prompt_section:
                kb_prompt_section += "\n\n" + kg_topic_context
            else:
                kb_prompt_section = kg_topic_context
            logger.info("KG found topic context for '%s'", topic)
    except Exception as e:
        logger.debug("KG topic context failed: %s", e)

    return kb_context, kb_prompt_section


def _validate_and_humanize(master: Any, topic: str, report: Any) -> None:
    """Run validation checks and humanize text on MasterContent in place."""
    from clawed.validation import check_self_contained, validate_alignment, validate_master_content

    mc_errors = validate_master_content(master, topic)
    for err in mc_errors:
        report.warnings.append(err)
        logger.warning("Validation: %s", err)

    _align_score, align_issues = validate_alignment(master)
    for issue in align_issues:
        report.warnings.append(issue)

    all_text = " ".join(s.content for s in master.direct_instruction)
    delegation = check_self_contained(all_text)
    for d in delegation:
        report.warnings.append(d)

    # Sanitize non-Latin characters
    _sanitize_master_content(master)

    # Humanize text
    try:
        from clawed.humanize import humanize
        for section in master.direct_instruction:
            if section.content:
                section.content = humanize(section.content)
        for entry in master.vocabulary:
            if entry.definition:
                entry.definition = humanize(entry.definition)
        for sq in master.exit_ticket:
            if hasattr(sq, "question") and sq.question:
                sq.question = humanize(sq.question)
    except Exception as e:
        logger.debug("Humanize pass failed: %s", e)


async def _compile_core_views(
    master: Any, images: dict[str, Any], output_dir: Any, config: Any,
) -> tuple[list[Any], list[str], list[str]]:
    """Compile teacher DOCX, student DOCX, and PPTX. Returns (files, side_effects, errors)."""
    generated_files: list[Any] = []
    side_effects: list[str] = []
    errors: list[str] = []

    try:
        from clawed.compile_teacher import compile_teacher_view
        teacher_path = await compile_teacher_view(master, images, output_dir)
        generated_files.append(teacher_path)
        side_effects.append(f"Teacher lesson plan DOCX: {teacher_path.name}")
    except Exception as e:
        logger.error("Teacher DOCX compile failed: %s", e)
        errors.append(f"Teacher DOCX failed: {e}")

    try:
        from clawed.compile_student import compile_student_view
        student_path = await compile_student_view(master, images, output_dir)
        generated_files.append(student_path)
        side_effects.append(f"Student packet DOCX: {student_path.name}")
    except Exception as e:
        logger.error("Student DOCX compile failed: %s", e)
        errors.append(f"Student packet failed: {e}")

    try:
        from clawed.compile_slides import compile_slides
        pptx_path = await compile_slides(master, images, output_dir)
        generated_files.append(pptx_path)
        side_effects.append(f"Slideshow PPTX: {pptx_path.name}")
    except Exception as e:
        logger.error("Slides compile failed: %s", e)
        errors.append(f"Slideshow PPTX failed: {e}")

    return generated_files, side_effects, errors


async def _run_auto_chain(
    master: Any, persona: Any, config: Any, output_dir: Any,
    subject: str, grade: str, topic: str,
    generated_files: list[Any], side_effects: list[str],
) -> None:
    """Run zero-touch auto-chain: differentiation, game, journey, research, standards."""
    safe_title = master.title.replace(" ", "_")[:40]

    # Auto-differentiate (IEP/504 + ELL + Gifted)
    if config and getattr(config, "auto_differentiate", True):
        try:
            from clawed.llm import LLMClient
            diff_llm = LLMClient(config=config)
            lesson_summary = (
                f"Title: {master.title}\nObjective: {master.objective}\n"
                f"Topic: {getattr(master, 'topic', master.title)}"
            )
            # The three differentiation variants are independent — generate them
            # concurrently instead of one-after-another (was ~3x the latency on a
            # path the teacher is actively waiting on).
            async def _make_diff(diff_type: str, diff_label: str) -> tuple[Any, str] | None:
                try:
                    diff_prompt = (
                        f"Generate {diff_label} for this lesson:\n\n{lesson_summary}\n\n"
                        "Provide specific, actionable modifications. Use bullet points. "
                        "Be concise (under 500 words)."
                    )
                    diff_text = await diff_llm.generate(diff_prompt)
                    if diff_text and len(diff_text) > 50:
                        from clawed.humanize import humanize
                        diff_text = humanize(diff_text)
                        diff_path = output_dir / f"{safe_title}_diff_{diff_type}.docx"
                        from docx import Document as DocxDocument
                        doc = DocxDocument()
                        doc.add_heading(f"{master.title} — {diff_label}", 0)
                        for line in diff_text.split("\n"):
                            if line.strip():
                                doc.add_paragraph(line.strip())
                        doc.save(str(diff_path))
                        return diff_path, f"{diff_label}: {diff_path.name}"
                except Exception as e:
                    logger.debug("Diff %s failed: %s", diff_type, e)
                return None

            diff_results = await asyncio.gather(*[
                _make_diff(dt, dl) for dt, dl in [
                    ("iep_504", "IEP/504 Accommodations"),
                    ("ell", "ELL Scaffolding"),
                    ("advanced", "Gifted Extensions"),
                ]
            ])
            for res in diff_results:
                if res:
                    generated_files.append(res[0])
                    side_effects.append(res[1])
        except Exception as e:
            logger.debug("Auto-differentiation failed: %s", e)

    # Auto-generate review game
    if config and getattr(config, "auto_game", True):
        try:
            from clawed.compile_game import compile_game
            game_path = await compile_game(master=master, persona=persona,
                                           output_dir=output_dir, config=config)
            if game_path and game_path.exists():
                generated_files.append(game_path)
                side_effects.append(f"Review game: {game_path.name}")
        except Exception as e:
            logger.debug("Auto-game failed: %s", e)

    # Auto-generate learning journey
    try:
        from clawed.compile_journey import compile_journey
        journey_path = await compile_journey(master=master, persona=persona, output_dir=output_dir)
        if journey_path and journey_path.exists():
            generated_files.append(journey_path)
            side_effects.append(f"Learning journey: {journey_path.name}")
    except Exception as e:
        logger.debug("Auto-journey failed: %s", e)

    # Auto-deep-research
    try:
        from clawed.deep_research import deep_research
        from clawed.llm import LLMClient
        research_llm = LLMClient(config=config)
        research_report = await deep_research(
            topic=getattr(master, "topic", master.title),
            subject=subject, grade=grade, llm=research_llm, max_sub_queries=3,
        )
        if research_report and len(research_report) > 100:
            research_path = output_dir / f"{safe_title}_research.md"
            research_path.write_text(research_report, encoding="utf-8")
            generated_files.append(research_path)
            side_effects.append(f"Research report: {research_path.name}")
    except Exception as e:
        logger.debug("Auto-research failed: %s", e)

    # Auto-lookup state standards
    if config and getattr(config, "auto_standards", True):
        try:
            state = ""
            if config.teacher_profile:
                state = config.teacher_profile.state
            if state:
                from clawed.state_standards import get_standards_context_for_prompt
                standards_ctx = get_standards_context_for_prompt(state, [subject], [grade])
                if standards_ctx:
                    logger.info("Standards aligned for %s/%s/%s", state, subject, grade)
        except Exception as e:
            logger.debug("Standards lookup failed: %s", e)


async def _run_quality_review(
    master: Any, config: Any, persona: Any,
    standards_list: list[Any], generated_files: list[Any], report: Any,
) -> float | None:
    """Run quality review and voice scoring. Returns voice_score or None."""
    voice_score = None
    try:
        from clawed.llm import LLMClient
        from clawed.quality import score_voice_match

        llm = LLMClient(config=config)
        master_json = master.model_dump_json(indent=2)[:3000]
        review = await llm.review_lesson_package(
            lesson_json=master_json, standards_present=bool(standards_list),
            has_handout=len(generated_files) >= 2, has_slideshow=len(generated_files) >= 3,
        )
        report.quality_review_passed = review.get("passed", False)
        report.quality_review_issues = review.get("issues", [])
        if not report.quality_review_passed:
            for issue in report.quality_review_issues:
                report.warnings.append(f"[REVIEW] {issue}")
        persona_ctx = persona.to_prompt_context() if persona else ""
        all_text = " ".join(s.content for s in master.direct_instruction)
        voice_score = await score_voice_match(all_text, persona_ctx, llm)
        report.voice_check_passed = voice_score >= 3.0
        if voice_score < 3.0:
            report.warnings.append(
                f"[{FailureCode.VOICE_MISMATCH}] Voice match score: {voice_score:.1f}/5.0"
            )
    except Exception as e:
        logger.warning("Quality review/voice check failed: %s", e)
        report.quality_review_passed = False
        report.quality_review_issues = [f"Review failed: {type(e).__name__}"]
    return voice_score


def _build_bundle_response(
    master: Any, generated_files: list[Any], errors: list[str],
    side_effects: list[str], kb_prompt_section: str, report: Any,
    standards_list: list[Any], voice_score: float | None,
) -> ToolResult:
    """Format the final ToolResult response for the bundle."""
    lines: list[str] = []

    if generated_files:
        lines.append(f"Complete teaching package for: **{master.title}**")
        lines.append(f"{len(generated_files)} files generated:\n")
        core = [f for f in generated_files if f.suffix in (".docx", ".pptx")
                and "diff_" not in f.name]
        diffs = [f for f in generated_files if "diff_" in f.name]
        extras = [f for f in generated_files if f.suffix in (".html", ".md") or f not in core + diffs]
        if core:
            lines.append("Core lesson:")
            for f in core:
                lines.append(f"  {f.name}")
        if diffs:
            lines.append("Differentiated:")
            for f in diffs:
                label = f.stem.split("diff_")[-1].upper().replace("_", "/")
                lines.append(f"  {f.name} ({label})")
        if extras:
            lines.append("Extras:")
            for f in extras:
                lines.append(f"  {f.name}")

        quality_parts: list[str] = []
        try:
            if voice_score and voice_score > 0:
                quality_parts.append(f"{voice_score:.1f}/5.0 voice match")
            if report.images_embedded:
                quality_parts.append(f"{report.images_embedded} images")
            if standards_list:
                short_codes = [s.split(":")[0].strip() for s in standards_list[:4]]
                std_str = ", ".join(short_codes)
                if len(standards_list) > 4:
                    std_str += f" (+{len(standards_list) - 4} more)"
                quality_parts.append(f"Standards: {std_str}")
        except Exception:
            logger.debug("operation_failed", exc_info=True)
        if quality_parts:
            lines.append(f"\nQuality: {' | '.join(quality_parts)}")
    else:
        lines.append(f"Failed to generate package for: {master.title}")
        for err in errors:
            lines.append(f"  - {err}")

    if kb_prompt_section:
        lines.append("\nReferenced your existing materials on this topic.")
    if report.warnings:
        lines.append("\nQuality notes:")
        for w in report.warnings[:5]:
            lines.append(f"  - {w}")

    return ToolResult(text="\n".join(lines), files=generated_files, side_effects=side_effects)


class GenerateLessonBundleTool:
    """Generate a COMPLETE teaching package: lesson plan + student handout + slideshow."""

    risk_level = "write_local"

    @staticmethod
    def approval_description(params: dict[str, Any]) -> str:
        topic = str(params.get("topic") or "lesson bundle").strip()
        return (
            "Generate and write a complete local teaching package for "
            f"{topic}: lesson plan, student handout, and slideshow."
        )

    def schema(self) -> dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "generate_lesson_bundle",
                "description": (
                    "Generate a COMPLETE teaching package for a topic: "
                    "a lesson plan (DOCX), a student handout (DOCX), and "
                    "a slideshow (PPTX). All three files are created at once. "
                    "IMPORTANT: Always call search_my_materials FIRST to find "
                    "the teacher's existing materials on this topic before "
                    "calling this tool."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "topic": {
                            "type": "string",
                            "description": "The lesson topic",
                        },
                        "grade": {
                            "type": "string",
                            "description": "Grade level (e.g. '8', 'K')",
                            "default": "8",
                        },
                        "subject": {
                            "type": "string",
                            "description": "Subject area",
                            "default": "General",
                        },
                        "activity_type": {
                            "type": "string",
                            "description": (
                                "Activity structure for the lesson"
                            ),
                            "enum": [
                                "jigsaw",
                                "socratic_seminar",
                                "document_analysis",
                                "debate",
                                "gallery_walk",
                                "station_rotation",
                                "direct_instruction",
                                "general",
                            ],
                            "default": "general",
                        },
                        "include_images": {
                            "type": "boolean",
                            "description": "Include academic images in the slideshow (default: true)",
                            "default": True,
                        },
                    },
                    "required": ["topic"],
                },
            },
        }

    async def execute(
        self, params: dict[str, Any], context: AgentContext
    ) -> ToolResult:
        from clawed.models import LessonBrief, TeacherPersona, UnitPlan
        from clawed.standards import get_standards_for_lesson

        topic = params["topic"]
        grade = params.get("grade", "8")
        subject = params.get("subject", "General")
        activity_type = params.get("activity_type", "general")
        include_images = params.get("include_images", True)

        context.notify_progress(
            f"Working on your lesson materials for \"{topic}\" now — "
            f"this usually takes 2-4 minutes. I'll send everything when it's ready!"
        )

        from clawed.generation_report import GenerationReport
        report = GenerationReport()

        # Load config & persona
        config = context.config
        persona = TeacherPersona()
        if context.persona:
            try:
                persona = TeacherPersona(**context.persona)
            except Exception as e:
                logger.warning("NLAH_FAILURE=%s: %s", FailureCode.PERSONA_PARSE_ERROR, e)
                report.warnings.append(f"[{FailureCode.PERSONA_PARSE_ERROR}] Could not parse persona: {e}")

        state = ""
        if config.teacher_profile and config.teacher_profile.state:
            state = config.teacher_profile.state

        standards_list = get_standards_for_lesson(subject=subject, grade=grade, state=state, topic=topic)

        # Search teacher's materials
        kb_context, kb_prompt_section = await _search_teacher_materials(context, topic, report)

        # Build UnitPlan
        description = f"Introduction to {topic}"
        if activity_type and activity_type != "general":
            description = f"{activity_type.replace('_', ' ').title()} lesson on {topic}"

        unit = UnitPlan(
            title=f"{topic} Unit", subject=subject, grade_level=grade, topic=topic,
            duration_weeks=1, overview=f"A lesson on {topic}.{kb_context}",
            standards=standards_list,
            daily_lessons=[LessonBrief(lesson_number=1, topic=topic, description=description)],
        )

        # Generate MasterContent
        from clawed.lesson import generate_master_content
        logger.info("Generating master content for '%s' (grade=%s, subject=%s)", topic, grade, subject)
        try:
            master = await generate_master_content(
                lesson_number=1, unit=unit, persona=persona, config=config,
                state=state, teacher_materials=kb_prompt_section,
            )
        except Exception as e:
            logger.error("NLAH_FAILURE=%s: %s", FailureCode.API_FAILURE, e)
            return ToolResult(text=f"[{FailureCode.API_FAILURE}] Failed to generate lesson: {type(e).__name__}")

        # Validate + humanize
        _validate_and_humanize(master, topic, report)
        # Narrate the phases — a full bundle takes a few minutes, so keep the
        # teacher seeing real progress instead of a silent spinner.
        context.notify_progress("Lesson content written — now adding images and primary sources…")

        # Fetch images
        images: dict[str, Path] = {}
        if include_images:
            try:
                from clawed.image_pipeline import fetch_all_images
                images = await fetch_all_images(master, config, teacher_id=context.teacher_id)
                report.images_embedded = len(images)
                logger.info("Fetched %d images", len(images))
            except Exception as e:
                logger.warning("Image fetch failed: %s", e)
                report.warnings.append(f"Image fetch failed: {e}")

        # Compile views
        context.notify_progress("Building the teacher plan, student handout, and slides…")
        output_dir = Path("~/clawed_output").expanduser().resolve()
        if config and hasattr(config, "output_dir") and config.output_dir:
            output_dir = Path(config.output_dir).expanduser().resolve()
        output_dir.mkdir(parents=True, exist_ok=True)

        generated_files, side_effects, errors = await _compile_core_views(master, images, output_dir, config)

        # Track generation
        if generated_files:
            try:
                from clawed.agent_core.quality import record_generation
                record_generation(
                    teacher_id=context.teacher_id, generation_type="lesson_bundle",
                    topic=topic, subject=subject, grade=grade,
                    scores={"format": activity_type, "files_generated": len(generated_files)},
                )
            except Exception as e:
                logger.debug("Lesson count tracking failed: %s", e)

        # Quality review + voice scoring
        voice_score = await _run_quality_review(
            master, config, persona, standards_list, generated_files, report,
        )

        # Auto-chain
        await _run_auto_chain(
            master, persona, config, output_dir, subject, grade, topic,
            generated_files, side_effects,
        )

        # Quality gate warning
        try:
            qt = getattr(config, "quality_threshold", 3.5)
            if voice_score and voice_score < qt and voice_score > 0:
                report.warnings.append(
                    f"Quality below threshold ({voice_score:.1f}/{qt}) — "
                    "consider regenerating with more specific instructions"
                )
        except Exception:
            logger.debug("operation_failed", exc_info=True)

        return _build_bundle_response(
            master, generated_files, errors, side_effects,
            kb_prompt_section, report, standards_list, voice_score,
        )
