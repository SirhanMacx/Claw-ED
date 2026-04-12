"""Student-view DOCX compiler for MasterContent.

Compiles a MasterContent object into a student-facing handout Word document.
Answer keys, teacher scripts, and station answer keys are omitted; guided
notes show prompts with blank lines instead of answers.
No LLM calls — pure mechanical compilation.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.shared import RGBColor

    from clawed.master_content import MasterContent

logger = logging.getLogger(__name__)

_BLANK = "_____________"


# ═══════════════════════════════════════════════════════════════════════════
# Reading level estimation (public utility)
# ═══════════════════════════════════════════════════════════════════════════


def _count_syllables(word: str) -> int:
    """Estimate syllable count for a single word using a vowel-group heuristic."""
    word = word.lower().strip()
    if not word:
        return 1
    # Remove trailing silent-e
    if word.endswith("e") and len(word) > 2:
        word = word[:-1]
    # Count vowel groups
    count = len(re.findall(r"[aeiouy]+", word))
    return max(count, 1)


def estimate_reading_level(text: str) -> str:
    """Estimate reading level using Flesch-Kincaid grade level formula.

    FK Grade = 0.39 * (words/sentences) + 11.8 * (syllables/words) - 15.59

    Returns one of:
        "Grade 5-6"  — simple vocabulary, short sentences
        "Grade 7-8"  — standard academic vocabulary
        "Grade 9-10" — complex vocabulary, longer sentences
        "Grade 11-12" — advanced academic language
        "College"    — specialized terminology
    """
    if not text or not text.strip():
        return "Grade 7-8"

    sentences = re.split(r"[.!?]+", text)
    sentences = [s.strip() for s in sentences if len(s.strip().split()) >= 3]
    if not sentences:
        return "Grade 7-8"

    words = re.findall(r"[a-zA-Z']+", text)
    if not words:
        return "Grade 7-8"

    total_words = len(words)
    total_sentences = len(sentences)
    total_syllables = sum(_count_syllables(w) for w in words)

    words_per_sentence = total_words / total_sentences
    syllables_per_word = total_syllables / total_words

    fk_grade = 0.39 * words_per_sentence + 11.8 * syllables_per_word - 15.59

    if fk_grade <= 6.5:
        return "Grade 5-6"
    elif fk_grade <= 8.5:
        return "Grade 7-8"
    elif fk_grade <= 10.5:
        return "Grade 9-10"
    elif fk_grade <= 12.5:
        return "Grade 11-12"
    else:
        return "College"


# ═══════════════════════════════════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════════════════════════════════


class _DocBuilder:
    """Thin wrapper holding a python-docx Document and theme state."""

    def __init__(self, doc: DocxDocument, primary_hex: str,
                 bg_light_hex: str, images: dict[str, Path]) -> None:
        from docx.shared import RGBColor
        self.doc = doc
        self.primary_hex = primary_hex
        self.bg_light_hex = bg_light_hex
        self.images = images
        self._RGBColor = RGBColor

    def hex_rgb(self, h: str) -> RGBColor:
        return self._RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def para(self, text: str, bold: bool = False, italic: bool = False,
             size_pt: int = 11, color: tuple[int, int, int] | None = None) -> None:
        from docx.shared import Pt
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = self._RGBColor(*color)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def embed_image(self, image_spec: str, width_inches: float = 4.5) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        if not image_spec:
            return
        path = self.images.get(image_spec)
        if path and Path(path).exists():
            try:
                self.doc.add_picture(str(path), width=Inches(width_inches))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                logger.debug("Could not embed image %r: %s", image_spec, exc)

    def shaded_cell(self, cell, fill_hex: str) -> None:
        from docx.oxml.ns import qn
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()  # noqa: N806
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill_hex,
        })
        tcPr.append(shd)

    def callout_box(self, title: str, items: list[str], fill_hex: str,
                    title_hex: str = "FFFFFF") -> None:
        """Render a colored callout box as a single-column table."""
        from docx.shared import Pt
        tbl = self.doc.add_table(rows=1 + len(items), cols=1)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells[0]
        self.shaded_cell(hdr, fill_hex)
        run = hdr.paragraphs[0].add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.hex_rgb(title_hex)
        run.font.name = "Calibri"
        for i, item in enumerate(items):
            cell = tbl.rows[i + 1].cells[0]
            self.shaded_cell(cell, self.bg_light_hex)
            run = cell.paragraphs[0].add_run(f"\u2022 {item}")
            run.font.size = Pt(10)
            run.font.name = "Calibri"
        self.doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════════════════
# Section builders
# ═══════════════════════════════════════════════════════════════════════════


def _build_header(b: _DocBuilder, master: MasterContent) -> None:
    """Title, metadata, reading level, and name/date fields."""
    b.doc.add_heading(master.title, level=0)

    # Estimate reading level from student-facing text
    reading_sample_parts = [master.topic, master.objective]
    for section in master.direct_instruction:
        reading_sample_parts.append(section.content)
    for ps in master.primary_sources:
        reading_sample_parts.append(ps.content_text)
    reading_sample = " ".join(reading_sample_parts)
    reading_level = estimate_reading_level(reading_sample)

    meta_lines = [
        f"Subject: {master.subject}  |  Grade: {master.grade_level}",
        f"Topic: {master.topic}",
        f"Objective: {master.objective}",
    ]
    for line in meta_lines:
        b.para(line)

    b.para(f"Reading Level: {reading_level}", italic=True, size_pt=9,
           color=(100, 100, 100))

    b.doc.add_paragraph("Name: ___________________________  Date: ___________  Period: _____")
    b.doc.add_paragraph("")


def _build_page_header(b: _DocBuilder, master: MasterContent) -> None:
    """Set up the page header with lesson title and student fields."""
    from docx.shared import Pt, RGBColor

    section = b.doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]

    title_run = header_para.add_run(master.title)
    title_run.font.size = Pt(8)
    title_run.font.color.rgb = RGBColor(128, 128, 128)
    title_run.font.name = "Calibri"

    header_para.add_run("\t\t")

    fields_run = header_para.add_run("Name: _________ Date: _________ Period: _____")
    fields_run.font.size = Pt(8)
    fields_run.font.color.rgb = RGBColor(128, 128, 128)
    fields_run.font.name = "Calibri"


def _build_vocabulary(b: _DocBuilder, master: MasterContent) -> None:
    """Vocabulary table with blank definitions for student fill-in."""
    from docx.shared import RGBColor

    if not master.vocabulary:
        return
    b.heading("Vocabulary")
    table = b.doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ["Term", "Definition", "Context Sentence"], strict=False):
        b.shaded_cell(cell, b.primary_hex)
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for entry in master.vocabulary:
        row = table.add_row().cells
        row[0].text = entry.term
        row[1].text = ""  # left blank for students
        row[2].text = ""  # left blank for students
        if entry.image_spec:
            b.embed_image(entry.image_spec, width_inches=1.5)

    b.doc.add_paragraph("")


def _build_do_now(b: _DocBuilder, master: MasterContent) -> None:
    """Do Now section with answer lines."""
    b.heading("Do Now")
    b.para(master.do_now.stimulus)
    if master.do_now.questions:
        for i, q in enumerate(master.do_now.questions, 1):
            b.para(f"{i}. {q}")
            b.para(_BLANK)
            b.para("")
    b.doc.add_paragraph("")


def _build_direct_instruction(b: _DocBuilder, master: MasterContent) -> None:
    """Direct instruction sections (no teacher script in student view)."""
    b.heading("Direct Instruction")
    for section in master.direct_instruction:
        b.heading(section.heading, level=2)
        b.para(section.content)
        if section.key_points:
            b.para("Key Points:", bold=True)
            for kp in section.key_points:
                p = b.doc.add_paragraph(style="List Bullet")
                p.add_run(kp)
        b.embed_image(section.image_spec)
        b.doc.add_paragraph("")


def _build_guided_notes(b: _DocBuilder, master: MasterContent) -> None:
    """Guided notes with blanks (no answers)."""
    if not master.guided_notes:
        return
    b.heading("Guided Notes")
    for note in master.guided_notes:
        b.para(note.prompt, bold=True)
        b.para(_BLANK)
        b.doc.add_paragraph("")


def _build_primary_sources(b: _DocBuilder, master: MasterContent) -> None:
    """Primary source documents with scaffolding questions and answer blanks."""
    if not master.primary_sources:
        return
    b.page_break()
    b.heading("Primary Sources")
    for ps in master.primary_sources:
        b.heading(ps.title, level=2)
        b.para(f"Type: {ps.source_type}  |  Attribution: {ps.attribution}")
        b.para(ps.content_text)
        if ps.scaffolding_questions:
            b.para("Questions:", bold=True)
            for sq in ps.scaffolding_questions:
                p = b.doc.add_paragraph(style="List Bullet")
                p.add_run(sq)
                b.doc.add_paragraph(_BLANK)
        b.embed_image(ps.image_spec)
        b.doc.add_paragraph("")


def _build_stations(b: _DocBuilder, master: MasterContent) -> None:
    """Learning stations with embedded source text (no answer key)."""
    if not master.stations:
        return
    b.heading("Learning Stations")
    _source_lookup = {ps.id: ps for ps in master.primary_sources} if master.primary_sources else {}
    for station in master.stations:
        b.heading(station.title, level=2)
        ref_source = _source_lookup.get(station.source_ref)
        if ref_source and ref_source.content_text:
            b.para(f"Document: {ref_source.title}", bold=True, italic=True,
                   size_pt=10, color=(80, 80, 80))
            if ref_source.attribution:
                b.para(f"\u2014 {ref_source.attribution}", italic=True, size_pt=9,
                       color=(100, 100, 100))
            b.para(ref_source.content_text)
            b.embed_image(ref_source.image_spec)
            b.doc.add_paragraph("")
        b.para(f"Task: {station.task}")
        b.para("Directions:", bold=True)
        b.para(station.student_directions)
        if ref_source and ref_source.scaffolding_questions:
            b.doc.add_paragraph("")
            for qi, sq in enumerate(ref_source.scaffolding_questions, 1):
                b.para(f"{qi}. {sq}")
                b.para(_BLANK)
        else:
            b.para(_BLANK)
        b.doc.add_paragraph("")


def _build_jigsaw(b: _DocBuilder, master: MasterContent) -> None:
    """Jigsaw graphic organizer and group sharing notes."""
    from docx.shared import RGBColor

    jigsaw = getattr(master, "jigsaw", None)
    if not (jigsaw and jigsaw.graphic_organizer):
        return

    b.heading("Jigsaw Notes")
    cols = [c.strip() for c in jigsaw.graphic_organizer.split("|") if c.strip()]
    if cols:
        num_rows = max(jigsaw.num_expert_groups, 3) + 1
        tbl = b.doc.add_table(rows=num_rows, cols=len(cols))
        tbl.style = "Table Grid"
        for j, col_name in enumerate(cols):
            cell = tbl.rows[0].cells[j]
            b.shaded_cell(cell, b.primary_hex)
            cell.text = col_name
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Group Sharing Notes Table
    if jigsaw.num_expert_groups and jigsaw.num_expert_groups >= 2:
        b.doc.add_paragraph("")
        b.heading("Group Sharing Notes", level=2)
        b.para(
            "Directions: As your group members share about their stations, "
            "take notes on the topics you did not study.",
            italic=True,
        )
        sharing_cols = [c.strip() for c in jigsaw.graphic_organizer.split("|") if c.strip()]
        if not sharing_cols:
            sharing_cols = ["Topic", "Key Findings", "My Notes"]
        share_rows = jigsaw.num_expert_groups + 1
        share_tbl = b.doc.add_table(rows=share_rows, cols=len(sharing_cols))
        share_tbl.style = "Table Grid"
        for j, col_name in enumerate(sharing_cols):
            cell = share_tbl.rows[0].cells[j]
            b.shaded_cell(cell, b.primary_hex)
            cell.text = col_name
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        b.doc.add_paragraph("")

    if jigsaw.debrief_question:
        b.doc.add_paragraph("")
        b.para(f"Debrief: {jigsaw.debrief_question}", bold=True)
        b.para(_BLANK)
    b.doc.add_paragraph("")


def _build_creative_activity(b: _DocBuilder, master: MasterContent) -> None:
    """Creative activity section."""
    creative = getattr(master, "creative_activity", None)
    if not (creative and creative.title):
        return
    b.heading(f"Activity: {creative.title}")
    if creative.scenario:
        b.para(creative.scenario)
    if creative.roles:
        b.para("Your Role:", bold=True)
        b.para(_BLANK)
    if creative.student_directions:
        b.para("Directions:", bold=True)
        b.para(creative.student_directions)
    if creative.deliverable:
        b.para(f"What you will create: {creative.deliverable}", bold=True)
        b.para(_BLANK)
        b.para(_BLANK)
    b.doc.add_paragraph("")


def _build_exit_ticket(b: _DocBuilder, master: MasterContent) -> None:
    """Exit ticket (stimulus + question, no answer)."""
    if not master.exit_ticket:
        return
    b.page_break()
    b.heading("Exit Ticket")
    fw = getattr(master.exit_ticket[0], "response_framework", "") if master.exit_ticket else ""
    if fw:
        b.para(f"REMEMBER {fw.upper()}!", bold=True, size_pt=13)
    else:
        b.para("Answer in complete sentences using evidence from the documents.", bold=True)
    b.doc.add_paragraph("")
    for i, sq in enumerate(master.exit_ticket, 1):
        b.para(f"Q{i}: {sq.stimulus}", bold=True)
        if sq.stimulus_image_spec:
            b.embed_image(sq.stimulus_image_spec)
        b.para(sq.question)
        starters = getattr(sq, "sentence_starters", [])
        if starters:
            b.para("Use these sentence starters:", italic=True)
            for starter in starters:
                b.para(f"  \u2022 {starter}", italic=True)
        b.para(_BLANK)
        b.doc.add_paragraph("")


def _build_differentiation(b: _DocBuilder, master: MasterContent) -> None:
    """Support strategies callout boxes."""
    diff = master.differentiation
    b.heading("Support Strategies")
    if diff.struggling:
        b.callout_box(
            "Extra Support",
            diff.struggling,
            fill_hex="D4A017",
            title_hex="FFFFFF",
        )
    if diff.advanced:
        b.callout_box(
            "Challenge Extension",
            diff.advanced,
            fill_hex="2B7A98",
            title_hex="FFFFFF",
        )
    if diff.ell:
        b.callout_box(
            "Language Support",
            diff.ell,
            fill_hex="2D8B4E",
            title_hex="FFFFFF",
        )


def _build_homework(b: _DocBuilder, master: MasterContent) -> None:
    """Homework section."""
    if not master.homework:
        return
    b.heading("Homework")
    b.para(master.homework)


# ═══════════════════════════════════════════════════════════════════════════
# Public API (unchanged signature)
# ═══════════════════════════════════════════════════════════════════════════


async def compile_student_view(
    master: MasterContent,
    images: dict[str, Path],
    output_dir: Path,
) -> Path:
    """Compile a student-facing DOCX from a MasterContent object.

    Guided notes show prompts with blank lines (not answers).  Teacher scripts,
    station answer keys, exit-ticket answers, and differentiation notes are all
    omitted — this is the print-ready student handout.

    Args:
        master: The MasterContent source-of-truth object.
        images: Mapping of image_spec strings to local file Paths.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document

    from clawed.export_theme import get_color_theme
    from clawed.io import safe_filename

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Subject-aware color theme
    theme = get_color_theme(master.subject)
    b = _DocBuilder(
        doc=doc,
        primary_hex=theme["primary"],
        bg_light_hex=theme["bg_light"],
        images=images,
    )

    # Page header
    _build_page_header(b, master)

    # Build all sections
    _build_header(b, master)
    _build_vocabulary(b, master)
    _build_do_now(b, master)
    _build_direct_instruction(b, master)
    _build_guided_notes(b, master)
    _build_primary_sources(b, master)
    _build_stations(b, master)
    _build_jigsaw(b, master)
    _build_creative_activity(b, master)
    _build_exit_ticket(b, master)
    _build_differentiation(b, master)
    _build_homework(b, master)

    # Save
    safe = safe_filename(master.title)
    out_path = output_dir / f"{safe}_student.docx"
    doc.save(str(out_path))
    logger.info("Student view saved to %s", out_path)
    return out_path
