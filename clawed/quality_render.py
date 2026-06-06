"""Answer-key leakage gate for student-facing documents.

Hard rule enforced here: *never* let the teacher answer key (or other
teacher-only content) reach the student handout.  Before a student-facing
document is delivered, scan its rendered text for accidentally-leaked answer
keys / teacher copy and flag any matches.

This is intentionally dependency-light: only ``python-docx`` (already a project
dependency) is used, and only to read text out of a ``.docx`` file.  When given
a string the scan runs with no third-party deps at all.  Unsupported / missing
input degrades gracefully (returns no findings) rather than raising.

Default posture is WARN, not block — callers decide whether to escalate via
:func:`assert_clean`.  The regex set is curated to catch real teacher-only
leakage while avoiding false positives on legitimate student text (e.g. the
bare word "key", or "answer the question").
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

__all__ = [
    "LEAKAGE_PATTERNS",
    "StudentLeakageError",
    "assert_clean",
    "extract_text",
    "scan_student_leakage",
]

# Number of characters of surrounding context to keep in each finding snippet.
_SNIPPET_RADIUS = 40
# Cap snippet length so a runaway match cannot bloat a report/log line.
_SNIPPET_MAX = 120


class StudentLeakageError(ValueError):
    """Raised by :func:`assert_clean` when leaked teacher content is found."""

    def __init__(self, findings: list[dict[str, Any]]) -> None:
        self.findings = findings
        patterns = ", ".join(sorted({f["pattern"] for f in findings}))
        super().__init__(
            f"Student document contains {len(findings)} answer-key/teacher-only "
            f"leak(s): {patterns}"
        )


# ── Curated leakage patterns ────────────────────────────────────────────────
# Each entry: (label, compiled-regex, severity).  All matching is
# case-insensitive.  Patterns are written to be *specific* — they target
# teacher-only phrasing and avoid generic student-facing words.
#
# Severity scale:
#   "critical" — an explicit answer key / correct-answer reveal.
#   "high"     — teacher-only copy, scripts, or scoring guidance.
def _p(pattern: str) -> re.Pattern[str]:
    return re.compile(pattern, re.IGNORECASE)


LEAKAGE_PATTERNS: list[tuple[str, re.Pattern[str], str]] = [
    # --- Explicit answer reveals (critical) ---
    # "answer key" as a phrase (word-bounded), not "answer the key question".
    ("answer key", _p(r"\banswer\s*key\b"), "critical"),
    # "Answer:" label — require it to read as a label (followed by content or
    # end of line), not the verb in "answer the question".  Disallow a
    # following "the/a/this/that/each/your/in/with" so prompts don't trip it.
    (
        "answer:",
        _p(r"\banswer\s*:(?!\s*(?:the|a|an|this|that|each|your|in|with|all|both)\b)"),
        "critical",
    ),
    # "Ans:" / "Ans." abbreviation used as an answer label.
    ("ans:", _p(r"\bans\s*[:.]\s*\S"), "critical"),
    # "correct answer" (the phrase teacher keys use).
    ("correct answer", _p(r"\bcorrect\s+answer\b"), "critical"),
    # "correct: A" style single-letter MC key.
    ("correct: <letter>", _p(r"\bcorrect\s*:\s*[A-D]\b"), "critical"),
    # "Key:" used as an answer-key label (followed by content). Avoids the bare
    # word "key" and "key term"/"key idea"/"key point" study language.
    (
        "key:",
        _p(r"\bkey\s*:(?!\s*(?:term|terms|idea|ideas|point|points|word|words|concept|concepts|vocabulary|question|questions)\b)\s*\S"),
        "high",
    ),
    # --- Teacher-only copy / scripts (high) ---
    ("teacher copy", _p(r"\bteacher\s*(?:copy|version|edition|key)\b"), "high"),
    ("teacher note", _p(r"\bteacher\s*notes?\b"), "high"),
    ("teacher only", _p(r"\bteacher[\s\-]*only\b"), "high"),
    # --- Scoring / grading guidance (high) ---
    ("mark scheme", _p(r"\bmark\s*scheme\b"), "high"),
    ("rubric points", _p(r"\brubric\s*points?\b"), "high"),
    ("sample response", _p(r"\bsample\s*responses?\b"), "high"),
    ("model answer", _p(r"\bmodel\s*answers?\b"), "high"),
]


# ── Text extraction ─────────────────────────────────────────────────────────


def _extract_docx_text(path: Path) -> str:
    """Return all readable text from a .docx (paragraphs + table cells).

    Degrades gracefully: if python-docx is unavailable or the file cannot be
    parsed, returns an empty string and logs at debug level.
    """
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - python-docx is a hard dep
        logger.debug("python-docx unavailable, cannot scan %s: %s", path, exc)
        return ""

    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.debug("Could not open %s as .docx: %s", path, exc)
        return ""

    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(source: str | Path) -> str:
    """Coerce *source* into scannable text.

    *source* may be:
      * a ``.docx`` file path (str or Path) — text is read via python-docx;
      * a path to any other existing file — read as UTF-8 text (best effort);
      * a raw string — returned as-is.

    Never raises on bad input — returns an empty string instead so the gate
    degrades gracefully.
    """
    # Path object → dispatch on suffix / existence.
    if isinstance(source, Path):
        if source.suffix.lower() == ".docx":
            return _extract_docx_text(source)
        try:
            if source.exists() and source.is_file():
                return source.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            logger.debug("Could not read %s as text: %s", source, exc)
        return ""

    # String: it might be a path to a real file, otherwise treat as raw text.
    if isinstance(source, str):
        if source.lower().endswith(".docx"):
            p = Path(source)
            if p.exists():
                return _extract_docx_text(p)
            return ""
        # Heuristic: a short single-line string that points at an existing file
        # is treated as a file path; anything else is raw content.
        if "\n" not in source and len(source) < 1024:
            try:
                p = Path(source)
                if p.exists() and p.is_file():
                    return p.read_text(encoding="utf-8", errors="replace")
            except (OSError, ValueError):
                pass
        return source

    # Anything else (None, bytes, etc.) → nothing to scan.
    logger.debug("Unsupported leakage-scan input type: %s", type(source).__name__)
    return ""


# ── Scan ────────────────────────────────────────────────────────────────────


def _snippet(text: str, start: int, end: int) -> str:
    """Return a trimmed, single-line context window around [start, end)."""
    lo = max(0, start - _SNIPPET_RADIUS)
    hi = min(len(text), end + _SNIPPET_RADIUS)
    window = text[lo:hi].replace("\n", " ").replace("\r", " ").strip()
    window = re.sub(r"\s+", " ", window)
    if len(window) > _SNIPPET_MAX:
        window = window[: _SNIPPET_MAX - 1].rstrip() + "…"
    prefix = "…" if lo > 0 else ""
    suffix = "…" if hi < len(text) else ""
    return f"{prefix}{window}{suffix}"


def scan_student_leakage(source: str | Path) -> list[dict[str, Any]]:
    """Scan a student-facing document for leaked answer-key / teacher content.

    Args:
        source: A ``.docx`` path (str or Path) whose text is extracted via
            python-docx, any other file path, or a raw string.

    Returns:
        A list of finding dicts, each ``{"pattern", "snippet", "severity"}``.
        An empty list means the document is clean.  Findings are de-duplicated
        on (pattern, lowercased snippet) so a repeated leak is reported once.
    """
    text = extract_text(source)
    if not text or not text.strip():
        return []

    findings: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for label, regex, severity in LEAKAGE_PATTERNS:
        for match in regex.finditer(text):
            snippet = _snippet(text, match.start(), match.end())
            key = (label, snippet.lower())
            if key in seen:
                continue
            seen.add(key)
            findings.append(
                {"pattern": label, "snippet": snippet, "severity": severity}
            )
    return findings


def assert_clean(source: str | Path) -> None:
    """Raise :class:`StudentLeakageError` if *source* leaks teacher content.

    Use this only where a hard block is desired.  The default delivery path
    warns rather than blocks (to avoid false-positive blocking); call this
    explicitly when you want leakage to be fatal.
    """
    findings = scan_student_leakage(source)
    if findings:
        raise StudentLeakageError(findings)
